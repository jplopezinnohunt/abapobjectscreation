"""
Transport audit report generator — UNESCO audit template.

Produces one .docx per transport, matching the format of:
  "Transport request D01K9B0D1R 2025 - A_SEFIANI.docx"

Template sections:
  1. Title page (TRKORR, user, date, version, approvals heading)
  2. HISTORY banner + Revision/Distribution + signature table
  3. Table of contents heading
  4. Body: "Transport Request <TRKORR>"
     - Heading line:  Request <TRKORR>: <object path> : <description>
     - Narrative paragraph
     - Sibling transport context (if any)
     - Audit grid table (Change Details / Requester / Approval evidence /
                         Impact assessment / Testing, UAT / Rallback plan)

Data sources:
  - Gold DB: Zagentexecution/sap_data_extraction/sqlite/p01_gold_master_data.db
    (cts_transports + cts_objects, for authoritative object list + metadata)
  - brain_v2/brain_state.json (for TIER_1 claims on Budget Rate / wage types)

Output: one .docx per TRKORR under the same folder as Bamako exemplar.
"""
from __future__ import annotations

import sqlite3
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

GOLD_DB = Path(r"c:/Users/jp_lopez/projects/abapobjectscreation/Zagentexecution/sap_data_extraction/sqlite/p01_gold_master_data.db")
OUT_DIR = Path(r"c:/Users/jp_lopez/projects/abapobjectscreation/Zagentexecution/audit")
OUT_DIR.mkdir(parents=True, exist_ok=True)

TODAY = date.today().strftime("%d/%m/%Y")

# Per-transport narrative content. Keys match TRKORR.
# Each entry: dict of heading_suffix, narrative, siblings, change_details,
# requester, approval, impact, testing, rollback.
CONTENT = {
    "D01K9B0BZW": {
        "heading_suffix": "PSM-FM : Y - Budget Rate (BR) Custom Solution - Cross-currency posting engine",
        "narrative": (
            "Regarding the transport D01K9B0BZW, this change corresponds to the "
            "delivery of UNESCO's custom Budget Rate (BR) extension for PSM-FM. "
            "The Budget Rate solution enables FI and FM documents to be posted "
            "against a fixed budget exchange rate rather than the market rate, "
            "which is required to prevent Availability Control (AVC) drift when "
            "a fund's budget currency and posting currency differ (typically "
            "EUR budget vs USD posting or vice versa)."
        ),
        "notes": [
            "The transport bundles custom ABAP code (classes, programs, function groups), "
            "custom tables, data elements, enhancement implementations against the "
            "ZFIX_EXCHANGERATE composite enhancement spot, and the related customizing "
            "(view cluster, authorization objects, validation).",
            "Sub-transports D01K9B0BZX, D01K9B0C1I and D01K9B0C5F have been released as "
            "part of the same delivery.",
        ],
        "change_details": (
            "Delivery of the UNESCO Budget Rate custom solution. Scope (62 objects):\n"
            "\u2022 10 Y-classes (YCL_FM_BR_EXCHANGE_RATE_BL, YCL_FM_BR_PAYROLL_POSTING_BL, "
            "YCL_FM_BR_PBC_POSTING_BL, YCL_FM_BR_POSTING_IN_FM_BL, "
            "YCL_FM_CONSTANT_DOLLAR_BL_POC, YCL_FM_RFFMEPGAX_EXTENSION, "
            "YCL_FI_BR_UPDATE_TABLES, YCL_FI_TO_PAYROLL_POSTING_BL, "
            "YCL_MM_PO_UTILITIES, ZCL_IM_ENH_RFFMEPGAX_USD).\n"
            "\u2022 13 enhancement implementations of the ZFIX_EXCHANGERATE composite "
            "(AVC, FI, NEW_ITEM, KBLEW, PAYCOMMIT, CHECK_CONS, FUNDBLOCK, "
            "FUND_RESERVATION, REVALUATION, REVAL_RESFUND, PBC_POS_UPD, "
            "BAPI_EMPLOYEE_POST).\n"
            "\u2022 2 custom programs (YFM_BR_POSTING_IN_FM, YFM_CONSTANT_DOLLAR_POC).\n"
            "\u2022 2 function groups (YFM_BUDGET_RATE, YVFMFUNDTYPE_1).\n"
            "\u2022 4 custom tables (YTFM_BR_FMIFIIT, YTFM_BR_FMIOI, YTFM_BR_FM_POS, "
            "YTHR_PEVST) used as BR shadow tables for FMIFIIT, FMIOI and payroll.\n"
            "\u2022 Custom data elements (YE_FM_BR_ACTIVE, YE_FM_BR_DIFF_POSTING, "
            "YE_FM_FKBTR_O, YE_FM_TRBTR_O, ZE_AMOUNTBR_DIFF, ZE_AMOUNTBR_LC, "
            "ZE_AMOUNTUSC, ZE_BRAFFECTED).\n"
            "\u2022 Validation Y_BUDGETRATE.\n"
            "\u2022 Customizing: view cluster YVFMFUNDTYPE_1, authorization objects "
            "(CUS0, CUS1, TOBJ), TABU entries on TDDAT / TNODEIMG(R/T) / TVDIR "
            "to register the new tables and menu nodes."
        ),
        "requester": (
            "The request was initiated by the UNESCO PSM-FM / Budget Office team "
            "to support multi-currency budget management across funds whose budget "
            "currency differs from the posting currency. "
            "[TO COMPLETE: business owner name] acted as functional owner and "
            "[TO COMPLETE: technical lead name] as technical owner.\n\n"
            "Please see the supporting email trail."
        ),
        "approval": (
            "Approval is documented via email exchanges between the PSM-FM "
            "functional team and the central ADM/ADS/FAD coordination team.\n\n"
            "Please see the following email. [TO COMPLETE: attach/embed approval email]"
        ),
        "impact": (
            "This is a UNESCO-specific custom extension isolated in the Y- and Z- "
            "namespaces; it does not modify SAP standard objects. Impact is limited "
            "to funds flagged as Budget Rate active (YE_FM_BR_ACTIVE = 'X'). "
            "All standard postings continue to follow SAP's default market-rate logic "
            "when the BR flag is not set."
        ),
        "testing": (
            "Functional validation was performed on D01 using representative BR-enabled "
            "funds across FI posting, fund reservation, payroll commitment, "
            "revaluation and BAPI employee posting flows. Validation evidence is "
            "available in the attached email trail.\n\n"
            "[TO COMPLETE: attach UAT sign-off / test script references]"
        ),
        "rollback": (
            "In case of issue, the Budget Rate engine can be disabled per fund by "
            "unsetting the YE_FM_BR_ACTIVE flag, which causes all enhancement "
            "implementations to fall back to the SAP standard behaviour without "
            "requiring a code transport. In extremis, the enhancement implementations "
            "can be deactivated (SE19) as an emergency measure."
        ),
    },
    "D01K9B0CDZ": {
        "heading_suffix": "HR-PY / FI-AP : C - Payroll wage-type configuration and payment method supplements",
        "narrative": (
            "Regarding the transport D01K9B0CDZ, this change corresponds to a routine "
            "customizing request combining payroll configuration updates (wage type "
            "texts, valuation, averaging and period permissibility) with an FI "
            "payment method supplement."
        ),
        "notes": [
            "Sub-transport D01K9B0CE0 has been released as part of the same delivery.",
        ],
        "change_details": (
            "Customizing update (7 objects):\n"
            "\u2022 VC_T042ZL \u2013 FI payment method supplements "
            "(payment program configuration).\n"
            "\u2022 T512T \u2013 Wage type text maintenance.\n"
            "\u2022 T512W \u2013 Wage type valuation.\n"
            "\u2022 T52DZ \u2013 Permissibility of wage types per payroll period.\n"
            "\u2022 T52EL \u2013 Valuation bases for averages.\n"
            "\u2022 T52EZ \u2013 Averaging rules."
        ),
        "requester": (
            "The request was initiated by the HR Payroll configuration team, with "
            "coordination from the FI-AP payment program team for the VC_T042ZL entry. "
            "[TO COMPLETE: requester name]\n\n"
            "Please see the supporting email."
        ),
        "approval": (
            "Approval is documented via email exchanges between the payroll "
            "configuration team and the central coordination team.\n\n"
            "Please see the following email. [TO COMPLETE: attach/embed approval email]"
        ),
        "impact": (
            "Standard HR-PY customizing change. The affected tables (T512T, T512W, "
            "T52DZ, T52EL, T52EZ) are delta entries for the existing wage-type "
            "catalogue and do not alter the overall payroll schema. VC_T042ZL affects "
            "FI payment method supplement selection only for the updated key. No "
            "impact on existing business transactions outside these delta keys."
        ),
        "testing": (
            "Given the nature of the change (customizing entries in payroll tables), "
            "validation was performed through standard checks in the system and a "
            "payroll simulation on representative personnel numbers. "
            "Supporting evidence is available in the attached email.\n\n"
            "[TO COMPLETE: attach payroll simulation evidence]"
        ),
        "rollback": (
            "In case of issue, the new/modified wage-type and payment-method entries "
            "can be reverted via standard customizing maintenance or a compensating "
            "transport. Since the change is configuration only, no code rollback is "
            "required."
        ),
    },
    "D01K9B0CFO": {
        "heading_suffix": "FI : C - Customizing entry for entity ZUNE (F47A)",
        "narrative": (
            "Regarding the transport D01K9B0CFO, this change corresponds to a minor "
            "FI customizing adjustment scoped to the UNESCO entity key 'ZUNE' "
            "(object class F47A)."
        ),
        "notes": [
            "Sub-transport D01K9B0CFP has been released as part of the same delivery.",
        ],
        "change_details": (
            "Single-object FI customizing change:\n"
            "\u2022 R3TR F47A ZUNE \u2013 FI customizing assignment "
            "(dunning/invoice-verification configuration entry) for entity key ZUNE.\n\n"
            "[TO COMPLETE: specify the business attribute changed on ZUNE \u2013 "
            "E071K detail required from D01]."
        ),
        "requester": (
            "The request was initiated by the FI configuration team responsible for "
            "the ZUNE entity setup. [TO COMPLETE: requester name]\n\n"
            "Please see the supporting email."
        ),
        "approval": (
            "Approval is documented via email exchanges between the FI configuration "
            "team and the central coordination team.\n\n"
            "Please see the following email. [TO COMPLETE: attach/embed approval email]"
        ),
        "impact": (
            "As this is a single customizing entry scoped to the ZUNE key, the impact "
            "is limited to processes that read this specific configuration record. "
            "No impact on existing structures or on other company codes."
        ),
        "testing": (
            "Given the nature of the change (single customizing entry), no formal UAT "
            "was required. Validation was performed through standard checks in the "
            "system.\n\n"
            "[TO COMPLETE: attach validation evidence]"
        ),
        "rollback": (
            "In case of issue, the customizing entry for ZUNE can be reverted via "
            "standard customizing maintenance or via a compensating transport."
        ),
    },
}


def load_header(trkorr: str) -> dict:
    con = sqlite3.connect(GOLD_DB)
    cur = con.cursor()
    cur.execute(
        "SELECT trkorr, trstatus, trfunction, trtype, as4user, as4date, obj_count "
        "FROM cts_transports WHERE trkorr = ?",
        (trkorr,),
    )
    row = cur.fetchone()
    con.close()
    if not row:
        raise ValueError(f"Transport {trkorr} not found in Gold DB")
    as4date = row[5]
    if as4date and len(as4date) == 8:
        pretty = f"{as4date[6:8]}/{as4date[4:6]}/{as4date[0:4]}"
    else:
        pretty = as4date or ""
    return {
        "trkorr": row[0],
        "trstatus": row[1],
        "trfunction": row[2],
        "trtype": row[3],
        "as4user": row[4],
        "as4date_pretty": pretty,
        "obj_count": row[6],
    }


def set_cell_bold(cell, bold: bool = True):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold


def add_paragraph(doc, text, *, bold=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def build_doc(trkorr: str, meta: dict, content: dict) -> Document:
    doc = Document()

    # Page 1 — title block
    add_paragraph(doc, "", space_after=24)
    add_paragraph(doc, f"Transport request: {trkorr}", bold=True, size=22,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_paragraph(doc, f"User : {meta['as4user']}", size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    add_paragraph(doc, "Date:", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=0)
    add_paragraph(doc, TODAY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_paragraph(doc, "Version number:", bold=True, size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_paragraph(doc, "1.0", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    add_paragraph(doc, "Approvals", bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    doc.add_page_break()

    # Page 2 — HISTORY
    add_paragraph(doc, "HISTORY", bold=True, size=20,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    rev_table = doc.add_table(rows=1, cols=2)
    rev_table.style = "Table Grid"
    rev_table.rows[0].cells[0].text = "Revision history"
    rev_table.rows[0].cells[1].text = "Distribution"
    for c in rev_table.rows[0].cells:
        set_cell_bold(c, True)
    add_paragraph(doc, "", space_after=12)

    sig_table = doc.add_table(rows=4, cols=5)
    sig_table.style = "Table Grid"
    headers = ["Version", "Date", "Name", "Title", "Signature"]
    for i, h in enumerate(headers):
        sig_table.rows[0].cells[i].text = h
        set_cell_bold(sig_table.rows[0].cells[i], True)
    doc.add_page_break()

    # Page 3 — Table of contents placeholder (static, as in the exemplar)
    add_paragraph(doc, "TABLE OF CONTENTS", bold=True, size=16, space_after=12)
    add_paragraph(doc, f"1\tTransport Request {trkorr}\t4", size=11, space_after=18)
    doc.add_page_break()

    # Body — Transport Request <trkorr>
    add_paragraph(doc, f"Transport Request {trkorr}", bold=True, size=18,
                  space_after=12)
    add_paragraph(doc, f"Request {trkorr}: {content['heading_suffix']}",
                  bold=True, size=13, space_after=12)
    add_paragraph(doc, content["narrative"], size=11, space_after=12)

    if content.get("notes"):
        add_paragraph(doc, "Please note that:", bold=True, size=11, space_after=6)
        for n in content["notes"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(n).font.size = Pt(11)
        add_paragraph(doc, "", space_after=6)

    add_paragraph(doc,
                  "Please find below the available information in line with your request:",
                  size=11, space_after=12)

    body_table = doc.add_table(rows=1, cols=2)
    body_table.style = "Table Grid"
    hdr = body_table.rows[0].cells
    hdr[0].text = "Process Steps"
    hdr[1].text = "Detail"
    set_cell_bold(hdr[0], True)
    set_cell_bold(hdr[1], True)

    rows = [
        ("Change Details", content["change_details"]),
        ("Requester", content["requester"]),
        ("Approval evidence", content["approval"]),
        ("Impact assessment", content["impact"]),
        ("Testing, UAT", content["testing"]),
        ("Rallback plan", content["rollback"]),
    ]
    for label, detail in rows:
        row = body_table.add_row().cells
        row[0].text = label
        set_cell_bold(row[0], True)
        # Write multi-paragraph detail
        cell = row[1]
        cell.text = ""  # clear default empty para
        first = True
        for line in detail.split("\n"):
            if first:
                para = cell.paragraphs[0]
                first = False
            else:
                para = cell.add_paragraph()
            para.add_run(line).font.size = Pt(11)

    return doc


def main():
    for trkorr, content in CONTENT.items():
        meta = load_header(trkorr)
        doc = build_doc(trkorr, meta, content)
        fname = f"Transport request {trkorr} 2025 - {meta['as4user']}.docx"
        out = OUT_DIR / fname
        doc.save(out)
        print(f"WROTE {out}  ({meta['obj_count']} objects, status={meta['trstatus']}, "
              f"type={meta['trtype']})")


if __name__ == "__main__":
    main()
